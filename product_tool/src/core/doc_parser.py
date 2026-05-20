# -*- coding: utf-8 -*-
"""
Word文档解析器 - 读取.docx文件
依赖: pip install python-docx
"""
import os
import re
import logging
from typing import Optional, List, Dict
from pathlib import Path

# python-docx (可选)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def read_docx(file_path: str) -> Optional[Document]:
    """
    读取Word文档
    
    Returns:
        Document对象 或 None
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed: pip install python-docx")
    
    if not os.path.exists(file_path):
        return None
    
    try:
        return Document(file_path)
    except Exception as e:
        logging.error(f'Error reading {file_path}: {e}')
        return None


def extract_text_from_docx(doc: Document) -> str:
    """
    提取纯文本
    
    Returns:
        str: 所有段落文本
    """
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def extract_tables(doc: Document) -> List[List[List[str]]]:
    """
    提取所有表格
    
    Returns:
        list: [[row1], [row2], ...], 每个row是[str]列表
    """
    tables = []
    
    for table in doc.tables:
        table_data = []
        prev_row = [None] * min(len(table.rows[0].cells) if table.rows else 0, 50)
        for ri, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            # 合并单元格传播：空值用前一行同列值填充（首行不传播也不记录，防表头泄漏）
            if ri > 0:
                for ci, val in enumerate(row_data):
                    if not val and ci < len(prev_row) and prev_row[ci]:
                        row_data[ci] = prev_row[ci]
            table_data.append(row_data)
            # 记录非空值作为下一行的合并参考（跳过首行）
            if ri > 0:
                for ci, val in enumerate(row_data):
                    if val and ci < len(prev_row):
                        prev_row[ci] = val
        
        if table_data:
            tables.append(table_data)
    
    return tables


def extract_text_blocks(doc: Document) -> List[Dict]:
    """
    提取文本块 (段落+表格)
    
    Returns:
        list: [{'type': 'paragraph'|'table', 'data': ...}]
    """
    blocks = []
    
    # 段落
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            blocks.append({'type': 'paragraph', 'data': text})
    
    # 表格
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        if table_data:
            blocks.append({'type': 'table', 'data': table_data})
    
    return blocks


# ==================== 内容分析辅助函数 ====================

def _infer_docx_columns(table: List[List[str]]) -> dict:
    """从DOCX表格数据内容推断每列的列角色（含包装列检测）"""
    if not table or len(table) < 2:
        return {}
    
    PACKAGING_KEYWORDS = {
        'gw', 'nw', 'g.w.', 'n.w.', 'gross weight', 'net weight',
        'gross', 'net', 'carton size', 'carton', 'package size',
        'packing size', 'cbm', 'meas', 'measurement',
        'qty/ctn', 'pcs/ctn', '每箱数量', '毛重', '净重', '外箱尺寸',
        '包装尺寸', '体积', 'cartons', '包装'
    }
    
    col_roles = {}
    for c in range(min(len(table[0]), 20)):
        texts = []
        price_count = 0
        model_count = 0
        label_count = 0
        packaging_count = 0
        
        for r in range(1, min(len(table), 50)):
            v = str(table[r][c] or '').strip()
            if not v:
                continue
            texts.append(v)
            v_lower = v.lower()
            if v_lower in {'image', 'picture', 'photo', 'description', 'name',
                           '备注', 'note', 'remark', '图片', '颜色', '颜色',
                           '尺寸', 'size', '规格', 'spec', '参数', '材质',
                           'material', '重量', 'weight', '功率', 'voltage'}:
                label_count += 1
            if v_lower in PACKAGING_KEYWORDS or any(kw in v_lower for kw in PACKAGING_KEYWORDS):
                packaging_count += 1
            if re.search(r'[A-Za-z]+\d+', v):
                model_count += 1
            # 尝试解析为价格：先清除单位后缀（/box, /pc, /pack, /set等）和货币符号
            cleaned_for_price = re.sub(r'/[a-zA-Z/]+', '', v)  # 去掉 /box, /pc, /pack 等
            cleaned_for_price = re.sub(r'[¥$€£,元]', '', cleaned_for_price)
            try:
                f = float(cleaned_for_price)
                if 0.01 < f < 10000000:
                    price_count += 1
            except Exception:
                pass
        
        if not texts:
            col_roles[c] = 'skip'
        elif packaging_count >= max(1, len(texts) * 0.15):
            col_roles[c] = 'packaging'
        elif label_count >= max(2, len(texts) * 0.3):
            col_roles[c] = 'label'
        elif price_count >= max(2, len(texts) * 0.2):
            col_roles[c] = 'price'
        elif model_count >= max(2, len(texts) * 0.15):
            col_roles[c] = 'model'
        else:
            avg_len = sum(len(t) for t in texts) / len(texts)
            col_roles[c] = 'model' if avg_len < 30 else 'spec'
    
    return col_roles


def _extract_products_from_text(text: str) -> List[Dict]:
    """从段落文本提取产品（适用于无表格的DOCX）"""
    products = []
    lines = text.split('\n')
    
    # 尝试按常见分隔模式提取
    current = {}
    for line in lines:
        line = line.strip()
        if not line:
            if current:
                products.append(current)
                current = {}
            continue
        
        # 检测是否新产品的起始行（含型号模式的行）
        has_model = bool(re.search(r'[A-Za-z]+\d+', line))
        has_price_line = bool(re.search(r'(?:价格|price|¥|\$)\s*[:：]?\s*[\d,.]+', line, re.I))
        
        if has_model and has_price_line and current:
            products.append(current)
            current = {'model': line}
        elif has_model and not has_price_line and not current.get('model'):
            current['model'] = line
        elif has_price_line:
            # 优先取 ¥/元/RMB 价格，没有才取 $
            m = re.search(r'(?:¥|元|RMB|人民币|cny)\s*([\d,.]+)', line, re.I)
            if m:
                current['price'] = m.group(1)
                current['currency'] = 'CNY'
            else:
                m = re.search(r'\$?\s*([\d,.]+)\s*(?:USD|美元|\$)?', line, re.I)
                if m:
                    current['price'] = m.group(1)
                    current['currency'] = 'USD'
                else:
                    m = re.search(r'(?:价格|price)\s*[:：]?\s*([\d,.]+)', line, re.I)
                    if m:
                        current['price'] = m.group(1)
        elif ':' in line or '：' in line:
            parts = re.split(r'[:：]', line, 1)
            key = parts[0].strip().lower()
            val = parts[1].strip()
            if any(kw in key for kw in ['型号', 'model', '产品', '品名', 'item']):
                if current.get('model'):
                    products.append(current)
                    current = {'model': val}
                else:
                    current['model'] = val
            elif any(kw in key for kw in ['价格', 'price', '单价']):
                current['price'] = val
            elif any(kw in key for kw in ['规格', 'spec', '参数', '尺寸']):
                current['spec'] = val
            elif any(kw in key for kw in ['名称', 'name', '品名']):
                current['name'] = val
    
    if current:
        products.append(current)
    
    return products


def _score_docx_result(products: List[Dict]) -> float:
    """评分DOCX解析结果（与 Excel score_result 公式一致）"""
    if not products:
        return -1
    n = len(products)
    has_model = sum(1 for p in products if p.get('model', '').strip())
    has_price = sum(1 for p in products if p.get('price'))
    model_set = set(str(p.get('model', '')).strip() for p in products if p.get('model'))
    diversity = len(model_set) / max(has_model, 1)
    score = n * 0.3 + has_model * 0.3 + has_price * 0.2 + diversity * 10 * 0.2
    # 产品_ 前缀惩罚（同 Excel）
    prefix_penalty = sum(1 for p in products if str(p.get('model', '')).startswith('产品_'))
    if prefix_penalty > n * 0.5:
        score *= 0.5
    return score


def _table_to_products(table, col_map, header_idx, currency='CNY'):
    """按列映射从表格提取产品"""
    products = []
    for row in table[header_idx + 1:]:
        if not any(row):
            continue
        product = {'currency': currency}
        for role, col_idx in col_map.items():
            if col_idx >= 0 and col_idx < len(row):
                product[role] = row[col_idx]
        # 如果 model 未找到且有 name，用 name 作为 model
        if not product.get('model') and product.get('name'):
            product['model'] = product['name']
        # model/name/spec 互检（原逻辑保留）
        model_val = str(product.get('model', ''))
        name_val = str(product.get('name', ''))
        if model_val and name_val:
            if (len(model_val) > len(name_val) * 1.5 or
                re.search(r'cm|mm|g|kg|inch|\d+\.?\d*\s*[xX×]\s*\d+', model_val, re.I) or
                re.search(r'\b[SMLXLXXL]\b', model_val, re.I) or   # S M L XL 等尺寸
                re.search(r'^\d+[\-/]', model_val)):               # 以数字开头如 '1.25cm...'
                product['model'] = name_val
                if not product.get('spec'):
                    product['spec'] = model_val
        # qty 空值时尝试用 Moq 列（如果映射了）
        if not product.get('qty') and col_map.get('qty', -1) >= 0:
            qty_raw = str(product.get('qty', '')).strip()
            import re as _re
            nums = _re.findall(r'\d+', qty_raw)
            if nums:
                product['qty'] = nums[0]
        if product:
            products.append(product)
    return products


def parse_product_docx(file_path: str) -> Optional[Dict]:
    """
    解析产品报价Word文档
    
    三策略择优:
    A. 表头关键词匹配（已有，增强）
    B. 自由文本提取（段落→产品）
    C. 内容推断列角色（兜底）
    
    Returns:
        dict: {
            'text': str,
            'tables': [...],
            'products': [{model, name, spec, price}, ...]
        }
    """
    doc = read_docx(file_path)
    if not doc:
        return None
    
    text = extract_text_from_docx(doc)
    tables = extract_tables(doc)
    
    all_candidates = []
    
    # ─── 策略A: 表头关键词匹配 ───
    try:
        from shared_keywords import DOCX_HEADER_KEYWORDS
        header_keywords = DOCX_HEADER_KEYWORDS
    except ImportError:
        header_keywords = ['型号', 'model', 'item', '产品', '名称', 'price', '价格',
                          '规格', 'spec', '参数', '数量', 'qty']
    
    for table in tables:
        if not table:
            continue
        
        header_idx = -1
        header_row = None
        for i, row in enumerate(table):
            row_text = ' '.join(row).lower()
            if any(kw in row_text for kw in header_keywords):
                header_idx = i
                header_row = row
                break
        
        if header_row:
            col_map = {'model': -1, 'name': -1, 'spec': -1, 'price': -1, 'qty': -1}
            for j, col in enumerate(header_row):
                col_lower = col.lower()
                if col_map['model'] < 0 and any(kw in col_lower for kw in ['型号', 'model', 'item', 'sku', '产品编号', '物料编码']):
                    # model↔spec 冲突：含强 spec 指示词时跳过
                    if any(kw in col_lower for kw in ['规格', 'spec', '参数', 'description']):
                        strong_spec = ['specification', 'specifications', '规格型号', '参数', 'description']
                        strong_model = ['型号', '物料编码', '料号', '货号', '产品编号', 'sku', 'item no', 'part no', 'p/n']
                        if any(kw in col_lower for kw in strong_spec) and not any(kw in col_lower for kw in strong_model):
                            pass  # 跳过 model 匹配
                        else:
                            col_map['model'] = j
                    else:
                        col_map['model'] = j
                elif col_map['name'] < 0 and any(kw in col_lower for kw in ['名称', 'name', '品名', '产品名称']):
                    col_map['name'] = j
                elif col_map['spec'] < 0 and any(kw in col_lower for kw in ['规格', 'spec', '参数', 'description', '描述', '尺寸']):
                    col_map['spec'] = j
                elif col_map['price'] < 0 and any(kw in col_lower for kw in ['价格', 'price', '单价', '金额', '元', 'rmb', 'cny', 'usd']):
                    col_map['price'] = j
                elif col_map['qty'] < 0 and any(kw in col_lower for kw in ['数量', 'qty', 'quantity', 'pcs', 'moq']):
                    col_map['qty'] = j
            
            if col_map['model'] >= 0 or col_map['price'] >= 0:
                # 检测币种：查看价格列表头文本
                detected_currency = 'CNY'
                if col_map['price'] >= 0 and col_map['price'] < len(header_row):
                    ph = header_row[col_map['price']].lower()
                    if 'usd' in ph or '$' in ph:
                        detected_currency = 'USD'
                products = _table_to_products(table, col_map, header_idx, currency=detected_currency)
                if products:
                    all_candidates.append(('table_header', products))
        
        # ─── 策略C: 内容推断（对每个表走一次） ───
        roles = _infer_docx_columns(table)
        model_c = None; price_c = None; spec_c = None; name_c = None; packaging_cols = []
        for c, role in sorted(roles.items()):
            if role == 'model' and model_c is None:
                model_c = c
            elif role == 'price' and price_c is None:
                price_c = c
            elif role == 'spec' and spec_c is None:
                spec_c = c
            elif role == 'packaging':
                packaging_cols.append(c)
        
        if model_c is not None:
            # 如果有包装列，把第一列包装列当 spec（若无 spec 列）
            effective_spec = spec_c if spec_c is not None else (packaging_cols[0] if packaging_cols else -1)
            cm = {'model': model_c, 'name': name_c if name_c is not None else model_c,
                  'spec': effective_spec, 'price': price_c if price_c is not None else -1, 'qty': -1}
            products = _table_to_products(table, cm, 0)
            if products:
                all_candidates.append(('table_content', products))
    
    # ─── 策略B: 段落文本提取 ───
    if text:
        text_products = _extract_products_from_text(text)
        if text_products:
            all_candidates.append(('paragraph', text_products))
    
    # ─── 策略D: 兜底 — 所有策略都空时，尝试内容推断无表头模式 ───
    if not all_candidates:
        for table in tables:
            if not table or len(table) < 2:
                continue
            roles = _infer_docx_columns(table)
            # 只要有疑似价格列或型号列，就提取
            has_price_col = any(r == 'price' for r in roles.values())
            has_model_col = any(r == 'model' for r in roles.values())
            if has_price_col or has_model_col:
                model_c = next((c for c, r in sorted(roles.items()) if r == 'model'), 0)
                price_c = next((c for c, r in sorted(roles.items()) if r == 'price'), -1)
                cm = {'model': model_c, 'name': model_c,
                      'spec': -1, 'price': price_c, 'qty': -1}
                products = _table_to_products(table, cm, 0)
                if products:
                    all_candidates.append(('table_fallback', products))
                    break
    
    # ─── 择优 ───
    if all_candidates:
        # 给表头匹配结果加分（正确识别列名比纯推断更可靠）
        scored = []
        for source, prods in all_candidates:
            s = _score_docx_result(prods)
            if source == 'table_header':
                s += 5  # 表头匹配优先
                # 额外加分：有价格列映射
                if any(p.get('price') for p in prods):
                    s += 3
            scored.append((s, source, prods))
        best = max(scored, key=lambda c: c[0])
        best_source = best[1]
        best_products = best[2]
        result = {
            'text': text,
            'tables': tables,
            'products': best_products,
            'parse_source': best_source,
        }
    else:
        result = {
            'text': text,
            'tables': tables,
            'products': [],
            'parse_source': 'empty',
        }
    
    return result


def extract_products_from_docx(file_path: str) -> 'pd.DataFrame':
    """从Word文档提取产品并转为DataFrame格式
    
    Returns:
        DataFrame with columns: model, name_zh, spec_zh, price_rmb, _row, _image_path
    """
    import pandas as pd
    
    doc_result = parse_product_docx(file_path)
    if not doc_result or not doc_result.get('products'):
        return pd.DataFrame()
    
    products = doc_result['products']
    data = []
    for idx, p in enumerate(products):
        price_val = 0
        price_raw_str = p.get('price', '')
        cur = p.get('currency', 'CNY')
        if cur not in ('USD', 'CNY'):
            cur = 'CNY'
        
        if price_raw_str:
            try:
                import re
                nums = re.findall(r'[\d,]+\.?\d*', str(price_raw_str))
                if nums:
                    price_val = float(nums[0].replace(',', ''))
            except (ValueError, TypeError):
                pass
        
        row = {
            'model': p.get('model', ''),
            'name_zh': p.get('name', ''),
            'spec_zh': p.get('spec', ''),
            'price_rmb': price_val,
            'price_raw': price_raw_str,  # 保留原始价格字符串（含单位/多价格）
            'currency': cur,
            'qty': p.get('qty', ''),
            '_row': idx + 1,
            '_sheet': os.path.basename(file_path),
        }
        data.append(row)
    
    df = pd.DataFrame(data)
    df = df[df['model'].notna() & (df['model'] != '')]
    
    # 图片匹配: DOCX 按顺序分配
    try:
        from .image import match_images_to_products_docx
        df = match_images_to_products_docx(df, file_path)
    except Exception:
        pass
    
    return df


def convert_docx_to_text(file_path: str, output_path: str = None) -> str:
    """
    转换Word为纯文本
    
    Args:
        file_path: .docx文件
        output_path: 输出.txt路径, None则返回文本
    
    Returns:
        str: 文本内容
    """
    doc = read_docx(file_path)
    if not doc:
        return ''
    
    text = extract_text_from_docx(doc)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        logging.info(f'Saved: {output_path}')
    
    return text


# ============ CLI ============

if __name__ == '__main__':
    import sys
    import glob
    
    if len(sys.argv) > 1:
        # convert <file.docx> [output.txt]
        file_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        
        if file_path.endswith('.docx'):
            text = convert_docx_to_text(file_path, output_path)
            if not output_path:
                print(text[:1000])  # 打印前1000字
        else:
            # 解析产品
            result = parse_product_docx(file_path)
            if result:
                print(f'Tables: {len(result["tables"])}')
                print(f'Products: {len(result["products"])}')
                print(result.get('text', '')[:1000])
            else:
                print('Parse failed')
    else:
        print("Word文档解析器")
        print("Usage: python doc_parser.py <file.docx> [output.txt]")
        print("Required: pip install python-docx")