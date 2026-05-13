"""
Document parser for product data extraction and normalization.

Supports multiple formats:
1. Excel (.xlsx, .xls) - Table format and Key-Value format
2. PDF (.pdf) - Extract tables from each page
3. DOCX (.docx) - Word documents with tables
"""
import os
import re
import glob
import shutil
from typing import Optional, Dict, List
import pandas as pd

# Try to import xlrd for old .xls files
try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False
    print("Note: xlrd not installed. Run 'pip install xlrd' to read old .xls files.")


# ==================== COLUMN CONFIG ====================

# 字段类型标记：core = 不参与序列化，spec = 参与序列化
COLUMN_CONFIG = {
    "model": {
        "keywords": ["型号", "Model", "产品编号", "SKU", "serial number", "no", "item no", "item", "code", "货号", "商品编号", "编号", "Model No.", "model number"],
        "type": "core"
    },
    "name_zh": {
        "keywords": ["产品名称", "名称", "中文名", "name", "品名", "商品名称", "产品名", "Product Name", "product name"],
        "type": "core"
    },
    "name_en": {
        "keywords": ["Product Name", "Name", "英文名", "English Name", "product name_en"],
        "type": "core"
    },
    "spec_zh": {
        "keywords": ["产品说明", "规格", "说明", "spec", "description", "产品描述", "描述", "参数", "specification", "product information", "product info"],
        "type": "spec"
    },
    "spec_en": {
        "keywords": ["specification", "specs", "description"],
        "type": "spec"
    },
    "color": {
        "keywords": ["颜色", "色", "color", "色彩"],
        "type": "spec"
    },
    "package": {
        "keywords": ["包装", "package", "包装规格", "规格包装"],
        "type": "spec"
    },
    "price_rmb": {
        "keywords": ["价格", "单价", "出厂价格", "RMB", "CNY", "价格(人民币)", "人民币", "价", "售价"],
        "type": "core"
    },
    "price_usd": {
        "keywords": ["Price", "USD", "$", "Unit price", "价格(美元)", "美金", "美元", "unit price", "Price($)"],
        "type": "core"
    },
    "quantity": {
        "keywords": ["数量", "Quantity", "MOQ", "最小起订量", "起订量"],
        "type": "core"
    },
    "image": {
        "keywords": ["图片", "Picture", "Photo", "图像", "产品图片", "产品照片", "image path", "image_path", "照片"],
        "type": "core"
    },
    "moq": {
        "keywords": ["MOQ", "最小起订量", "起订量", "最小订量"],
        "type": "spec"
    },
    "lead_time": {
        "keywords": ["lead time", "交货期", "leadtime", "交货时间"],
        "type": "spec"
    },
    "certification": {
        "keywords": ["certification", "认证", "证书", "cert"],
        "type": "spec"
    }
}

# 字段别名映射：标准化字段名
FIELD_ALIASES = {
    "product name": "name_zh",
    "product name_en": "name_en",
    "name": "name_zh",
    "spec": "spec_zh",
    "specification": "spec_zh",
    "price": "price_rmb",
    "price rmb": "price_rmb",
    "price usd": "price_usd",
}

# 分隔符定义：用于拆分中英文对照文本
BILINGUAL_SEPARATORS = ["<br>", "/", "；", ";"]


def is_core_field(col_name: str) -> bool:
    """判断列是否为核心字段（不参与序列化）"""
    if not col_name:
        return False
    col_lower = col_name.lower().strip()
    for field_name, config in COLUMN_CONFIG.items():
        if config.get("type") == "core":
            for kw in config.get("keywords", []):
                if kw.lower() in col_lower or col_lower in kw.lower():
                    return True
    return False


def find_field_from_config(target_field: str) -> Optional[str]:
    """从 COLUMN_CONFIG 中查找目标字段的标准名"""
    if not target_field:
        return None
    target_lower = target_field.lower().strip()
    
    # 直接匹配
    if target_field in COLUMN_CONFIG:
        return target_field
    
    # 别名匹配
    for alias, std_name in FIELD_ALIASES.items():
        if alias.lower() == target_lower:
            return std_name
    
    # 关键词匹配
    for field_name, config in COLUMN_CONFIG.items():
        for kw in config.get("keywords", []):
            if kw.lower() == target_lower:
                return field_name
    return None


def get_field_keywords(field_name: str) -> list:
    """获取字段的关键词列表"""
    if field_name in COLUMN_CONFIG:
        return COLUMN_CONFIG[field_name].get("keywords", [])
    if field_name in FIELD_ALIASES:
        std_name = FIELD_ALIASES[field_name]
        return COLUMN_CONFIG.get(std_name, {}).get("keywords", [])
    return []


def get_spec_columns(columns: list) -> list:
    """获取所有规格参数列（排除核心字段）"""
    return [col for col in columns if not is_core_field(col)]


# ==================== SPEC SERIALIZATION FUNCTIONS ====================

def split_bilingual_text(text: str, lang: str = "zh") -> str:
    """
    拆分中英文对照文本
    
    支持的分隔符: <br>, /, ；, ;
    
    Args:
        text: 待拆分的文本
        lang: "zh" 取中文部分, "en" 取英文部分
    
    Returns:
        拆分后的文本
    """
    if not text or pd.isna(text):
        return ""
    
    text_str = str(text).strip()
    if not text_str:
        return ""
    
    # 尝试找到分隔符
    separator = None
    for sep in BILINGUAL_SEPARATORS:
        if sep in text_str:
            separator = sep
            break
    
    if not separator:
        return text_str
    
    parts = text_str.split(separator)
    
    if len(parts) == 1:
        return text_str
    
    if lang == "zh":
        # 取第一部分（中文）
        result = parts[0].strip()
        return result
    else:
        # 取最后一部分（英文）
        result = parts[-1].strip()
        return result


def clean_column_name(col_name: str, lang: str = "zh") -> str:
    """
    清理列名，移除多余的单位和注释
    
    Args:
        col_name: 列名
        lang: "zh" 取中文部分, "en" 取英文部分
    
    Returns:
        清理后的列名
    """
    if not col_name or pd.isna(col_name):
        return ""
    
    col_str = str(col_name).strip()
    if not col_str:
        return ""
    
    # 先拆分中英文对照
    text = split_bilingual_text(col_str, lang)
    
    # 移除常见的单位注释
    patterns_to_remove = [
        r'\(CNY\)', r'\(RMB\)', r'\(USD\)', r'\(\$\)',
        r'\(未含电池\)', r'\(含电池\)', r'\(不含电池\)',
        r'\(kW\)', r'\(KW\)', r'\(W\)', r'\(V\)', r'\(A\)',
        r'\(kg\)', r'\(KG\)', r'\(L\)', r'\(mm\)',
    ]
    
    for pattern in patterns_to_remove:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # 清理多余空格
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def serialize_specs(df: pd.DataFrame, core_columns: list = None) -> pd.DataFrame:
    """
    将规格参数列序列化为 spec_text_zh 和 spec_text_en
    
    处理流程:
    1. 识别核心字段（不参与序列化）
    2. 识别规格参数列（其他所有列）
    3. 对每个规格参数:
       - 列名清理（中英文拆分）
       - 值清理（中英文拆分、空值过滤）
    4. 生成 spec_text_zh 和 spec_text_en
    
    Args:
        df: 原始 DataFrame
        core_columns: 可选，显式指定的核心字段列表
    
    Returns:
        添加了 spec_text_zh 和 spec_text_en 列的 DataFrame
    """
    if df is None or df.empty:
        return df
    
    result = df.copy()
    
    # 直接使用 is_core_field 判断每个原始列是否为核心字段
    if core_columns is None:
        core_columns = []
        for col in result.columns:
            if is_core_field(col):
                core_columns.append(col)
    
    # 获取规格参数列（排除核心字段和内部字段）
    internal_fields = ['_source_file', '_source_sheet', 'spec_dict', 'spec_zh', 'spec_dict']
    spec_columns = []
    for col in result.columns:
        if col in core_columns:
            continue
        if col in internal_fields:
            continue
        # 跳过已经是spec_text的列
        if col.startswith('spec_text_'):
            continue
        spec_columns.append(col)
    
    # 生成规格文本
    spec_text_zh_list = []
    spec_text_en_list = []
    
    for idx in range(len(result)):
        spec_lines_zh = []
        spec_lines_en = []
        
        for col in spec_columns:
            # 获取列名（清理后）
            col_name_zh = clean_column_name(col, "zh")
            col_name_en = clean_column_name(col, "en")
            
            # 获取值
            value = result.iloc[idx][col] if col in result.columns else None
            
            # 跳过空值
            if pd.isna(value) or str(value).strip() == "":
                continue
            
            value_str = str(value).strip()
            
            # 拆分中英文值
            value_zh = split_bilingual_text(value_str, "zh")
            value_en = split_bilingual_text(value_str, "en")
            
            # 跳过空值
            if not value_zh or value_zh == "nan":
                continue
            
            # 添加到规格列表
            if col_name_zh and value_zh:
                spec_lines_zh.append(f"{col_name_zh}: {value_zh}")
            
            if col_name_en and value_en:
                spec_lines_en.append(f"{col_name_en}: {value_en}")
            elif col_name_en and value_zh:
                # 如果英文值为空，使用中文值
                spec_lines_en.append(f"{col_name_en}: {value_zh}")
        
        spec_text_zh_list.append("\n".join(spec_lines_zh))
        spec_text_en_list.append("\n".join(spec_lines_en))
    
    # 添加新列
    result["spec_text_zh"] = spec_text_zh_list
    result["spec_text_en"] = spec_text_en_list
    
    return result


# Supported image extensions
IMAGE_EXTENSIONS = [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"]

# Temp folder for images
TEMP_IMAGE_FOLDER = "temp_images"

# Minimum rows to be considered a product sheet
MIN_PRODUCT_ROWS = 5

# Default exchange rate
DEFAULT_EXCHANGE_RATE = 7.2

# Global exchange rate variables
_current_exchange_rate = DEFAULT_EXCHANGE_RATE
_exchange_rate_date = None


def get_exchange_rate(base: str = "CNY", target: str = "USD") -> float:
    """Get real-time exchange rate."""
    global _current_exchange_rate, _exchange_rate_date
    
    import datetime
    import urllib.request
    import json
    
    apis = [
        f"https://api.exchangerate-api.com/v4/latest/{base}",
        f"https://open.er-api.com/v6/latest/{base}",
    ]
    
    for api_url in apis:
        try:
            req = urllib.request.Request(api_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=5) as response:
                data = json.loads(response.read().decode('utf-8'))
            
            if 'rates' in data and target in data['rates']:
                rate = data['rates'][target]
                _current_exchange_rate = rate
                _exchange_rate_date = datetime.datetime.now().strftime("%Y-%m-%d")
                print(f"  Rate: 1 {base} = {rate:.4f} {target} ({_exchange_rate_date})")
                return rate
        except Exception:
            continue
    
    print(f"  Warning: Using default rate: {DEFAULT_EXCHANGE_RATE}")
    _current_exchange_rate = DEFAULT_EXCHANGE_RATE
    _exchange_rate_date = datetime.datetime.now().strftime("%Y-%m-%d")
    return DEFAULT_EXCHANGE_RATE


def get_current_exchange_rate() -> float:
    """Get current exchange rate"""
    return _current_exchange_rate


def get_exchange_rate_date() -> str:
    """Get exchange rate date"""
    return _exchange_rate_date or datetime.datetime.now().strftime("%Y-%m-%d")


# Keywords to detect key-value format
MODEL_KEYWORDS = ["model:", "型号:", "型号：", "型号", "model"]
PRICE_KEYWORDS = ["price:", "price", "价格:", "价格：", "rmb", "price"]


# ==================== HELPER FUNCTIONS ====================

def _is_key_value_format(df: pd.DataFrame) -> bool:
    """Check if DataFrame uses key-value format."""
    if df is None or df.empty or df.shape[1] < 2:
        return False
    for i in range(min(20, len(df))):
        val = str(df.iloc[i, 1]).lower() if pd.notna(df.iloc[i, 1]) else ""
        for kw in MODEL_KEYWORDS:
            if kw.lower() in val:
                return True
    return False


# ==================== ENCODING FIX ====================

def fix_encoding(text) -> str:
    """Fix common encoding issues in text."""
    if not text or pd.isna(text):
        return ""
    
    text_str = str(text)
    
    if len(text_str) < 2:
        return text_str
    
    # Normalize full-width punctuation
    text_str = text_str.replace('\uff1a', ':')
    text_str = text_str.replace('\uff08', '(')
    text_str = text_str.replace('\uff09', ')')
    text_str = text_str.replace('\uff0d', '-')
    text_str = text_str.replace('\uff02', '"')
    text_str = text_str.replace('\uff07', "'")
    
    return text_str


def _find_column(df_columns: list, target_field: str) -> Optional[str]:
    """Find the best matching column name for a target field using COLUMN_CONFIG."""
    keywords = get_field_keywords(target_field)
    for col in df_columns:
        if not isinstance(col, str):
            continue
        col_lower = col.lower().strip()
        for kw in keywords:
            if kw.lower() in col_lower or col_lower in kw.lower():
                return col
    return None


# ==================== IMAGE HANDLING ====================

def _ensure_temp_image_folder() -> str:
    """Create temp image folder path."""
    folder = os.path.abspath(TEMP_IMAGE_FOLDER)
    os.makedirs(folder, exist_ok=True)
    return folder


def _copy_image_to_temp(source_path: str, model: str, output_folder: str = None) -> Optional[str]:
    """Copy image to temp folder with model-based naming."""
    if not source_path or pd.isna(source_path):
        return None
    
    source_path = str(source_path).strip()
    if not os.path.exists(source_path):
        return None
    
    _, ext = os.path.splitext(source_path.lower())
    if ext not in IMAGE_EXTENSIONS:
        return None
    
    folder = output_folder or _ensure_temp_image_folder()
    new_name = f"{model}{ext}"
    new_path = os.path.join(folder, new_name)
    
    try:
        shutil.copy2(source_path, new_path)
        return new_path
    except Exception:
        return None


def scan_for_images(source_folder: str) -> dict:
    """Scan folder for product images."""
    image_map = {}
    
    if not os.path.isdir(source_folder):
        return image_map
    
    for ext in IMAGE_EXTENSIONS:
        for img_path in glob.glob(os.path.join(source_folder, f"*{ext}")):
            basename = os.path.basename(img_path)
            model = os.path.splitext(basename)[0]
            model = re.sub(r'[_-]?(img|photo|product|pic)?$', '', model, flags=re.IGNORECASE)
            if model:
                image_map[model] = img_path
    
    return image_map


def match_images(df: pd.DataFrame, image_folder: str = "./images") -> pd.DataFrame:
    """Match images to products based on model numbers."""
    if df is None or df.empty:
        return df
    
    if "model" not in df.columns:
        print("  Warning: No 'model' column, cannot match images")
        return df
    
    result = df.copy()
    result["image_path"] = None
    
    image_map = scan_for_images(image_folder)
    
    if not image_map:
        print(f"  No images found in {image_folder}")
        return result
    
    print(f"  Found {len(image_map)} images, matching to products...")
    
    matched_count = 0
    for idx in result.index:
        model = result.loc[idx, "model"]
        if pd.isna(model) or not model:
            continue
        
        model_str = str(model).strip()
        
        if model_str in image_map:
            result.loc[idx, "image_path"] = image_map[model_str]
            matched_count += 1
        else:
            model_lower = model_str.lower()
            for img_model, img_path in image_map.items():
                if img_model.lower() == model_lower:
                    result.loc[idx, "image_path"] = img_path
                    matched_count += 1
                    break
    
    print(f"  Matched {matched_count} products to images")
    return result


# ==================== NORMALIZE DATAFRAME ====================

def _normalize_dataframe(df: pd.DataFrame, exchange_rate: float = DEFAULT_EXCHANGE_RATE) -> Optional[pd.DataFrame]:
    """Normalize a DataFrame to standard columns."""
    if df is None or df.empty:
        return None

    # 保留已存在的 spec_text 列
    spec_text_zh = df["spec_text_zh"] if "spec_text_zh" in df.columns else None
    spec_text_en = df["spec_text_en"] if "spec_text_en" in df.columns else None

    # Find relevant columns
    model_col = _find_column(df.columns.tolist(), "model")
    name_col = _find_column(df.columns.tolist(), "name_zh")
    spec_col = _find_column(df.columns.tolist(), "spec_zh")
    color_col = _find_column(df.columns.tolist(), "color")
    package_col = _find_column(df.columns.tolist(), "package")
    price_rmb_col = _find_column(df.columns.tolist(), "price_rmb")
    price_usd_col = _find_column(df.columns.tolist(), "price_usd")
    image_col = _find_column(df.columns.tolist(), "image")

    # Must have at least model or name
    if not model_col and not name_col:
        return None

    # Build normalized dataframe
    normalized_data = {}
    if model_col:
        normalized_data["model"] = df[model_col].apply(fix_encoding)
    if name_col:
        normalized_data["name_zh"] = df[name_col].apply(fix_encoding)
    if spec_col:
        def parse_spec_to_dict(spec_val):
            if pd.isna(spec_val) or not spec_val:
                return {}
            spec_str = str(spec_val).strip()
            if not spec_str:
                return {}
            
            spec_str = fix_encoding(spec_str)
            
            spec_dict = {}
            lines = spec_str.replace("\r\n", "\n").replace("\r", "\n").split("\n")
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if ":" in line:
                    parts = line.split(":", 1)
                    key = parts[0].strip()
                    value = parts[1].strip() if len(parts) > 1 else ""
                elif "-" in line:
                    parts = line.split("-", 1)
                    key = parts[0].strip()
                    value = parts[1].strip() if len(parts) > 1 else ""
                else:
                    continue
                
                key = fix_encoding(key)
                value = fix_encoding(value)
                
                if key and key.lower() not in ["型号", "model", "价格", "price", "name", "名称"]:
                    spec_dict[key] = value
            
            return spec_dict if spec_dict else {}

        normalized_data["spec_dict"] = df[spec_col].apply(parse_spec_to_dict)
        
        def dict_to_string(spec_dict):
            if not spec_dict or isinstance(spec_dict, str):
                return spec_dict if spec_dict else ""
            if isinstance(spec_dict, dict):
                return "; ".join(f"{k}: {v}" for k, v in spec_dict.items())
            return str(spec_dict) if spec_dict else ""
        
        normalized_data["spec_zh"] = df[spec_col].apply(dict_to_string)
        
    if color_col:
        normalized_data["color"] = df[color_col]
    if package_col:
        normalized_data["package"] = df[package_col]
    if price_rmb_col:
        normalized_data["price_rmb"] = df[price_rmb_col]
    if price_usd_col:
        normalized_data["price_usd"] = df[price_usd_col]
    
    # Handle image column
    if image_col:
        temp_folder = _ensure_temp_image_folder()
        def copy_image(img_path, model):
            if pd.isna(img_path) or not img_path:
                return None
            return _copy_image_to_temp(str(img_path), str(model), temp_folder)
        
        normalized_data["image_path"] = df.apply(
            lambda row: copy_image(row[image_col], row.get(model_col, "")), axis=1
        )

    result = pd.DataFrame(normalized_data)

    # Ensure all required columns exist
    for col in ["model", "name_zh", "spec_zh", "spec_dict", "color", "package", "price_rmb", "price_usd", "image_path"]:
        if col not in result.columns:
            result[col] = {} if col == "spec_dict" else None
    
    # 保留spec_text列
    if spec_text_zh is not None:
        result["spec_text_zh"] = spec_text_zh
    if spec_text_en is not None:
        result["spec_text_en"] = spec_text_en

    # ========== FIX: Data Cleaning ==========
    # Remove rows with empty model AND empty name_zh
    result = result[result["model"].notna() | result["name_zh"].notna()]
    
    # Remove header-like rows
    result = result[~result["model"].astype(str).str.lower().str.contains(
        "型号|name|model|价格|price", na=False, regex=True
    )]
    
    # Remove rows where name_zh is just "产品名称" or similar headers
    header_patterns = ["产品名称", "产品名称 ", "备注", "规格", "型号"]
    for pattern in header_patterns:
        mask = result["name_zh"].astype(str).str.strip() == pattern
        # Only remove if other columns are also empty
        other_cols = [c for c in result.columns if c not in ["name_zh", "model", "spec_dict"]]
        if other_cols:
            all_empty = result[other_cols].isna().all(axis=1)
            result = result[~(mask & all_empty)]
        else:
            result = result[~mask]
    
    result = result.reset_index(drop=True)
    
    return result


# ==================== EXCEL PARSING ====================

def _parse_key_value_format(df: pd.DataFrame, exchange_rate: float = DEFAULT_EXCHANGE_RATE) -> Optional[pd.DataFrame]:
    """Parse key-value format Excel."""
    if df is None or df.empty:
        return None

    products = []
    current_product = {}

    for idx in range(len(df)):
        key = str(df.iloc[idx, 1]).strip() if pd.notna(df.iloc[idx, 1]) else ""
        value = str(df.iloc[idx, 2]).strip() if pd.notna(df.iloc[idx, 2]) else ""

        key_clean = key.lower().replace("：", ":").replace("﻿", "").replace("\n", "")

        if any(kw.lower().replace(":", "") in key_clean.replace(":", "") for kw in MODEL_KEYWORDS):
            model_match = re.sub(r".*?:", "", key, flags=re.IGNORECASE).strip()
            model_name = value or model_match
            if current_product.get("model") == model_name:
                continue
            if current_product and (current_product.get("model") or current_product.get("name_zh")):
                if current_product.get("spec_dict"):
                    current_product["spec_zh"] = "; ".join(
                        f"{k}: {v}" for k, v in current_product["spec_dict"].items()
                    )
                products.append(current_product)
            current_product = {"model": model_name, "name_zh": model_name, "spec_dict": {}}

        price_col = df.iloc[idx, 3] if df.shape[1] > 3 else None
        if price_col and pd.notna(price_col):
            price_str = str(price_col).upper()
            price_match = re.search(r"rmb(\d[\d,]+\.?\d*)", price_str, re.IGNORECASE)
            if price_match:
                try:
                    price_val = float(price_match.group(1).replace(",", ""))
                    current_product["price_rmb"] = price_val
                    current_product["price_usd"] = round(price_val / exchange_rate, 2)
                except (ValueError, TypeError):
                    pass
            else:
                try:
                    price_val = float(re.sub(r"[^\d.]", "", str(price_col)))
                    if 50 < price_val < 100000:
                        current_product["price_rmb"] = price_val
                        current_product["price_usd"] = round(price_val / exchange_rate, 2)
                except (ValueError, TypeError):
                    pass
        else:
            if key and value and current_product:
                spec_key = re.sub(r".*?:", "", key).strip()
                spec_key = fix_encoding(spec_key)
                value = fix_encoding(value)
                if spec_key and spec_key not in ["model", "型号", "price", "价格"]:
                    if "spec_dict" not in current_product:
                        current_product["spec_dict"] = {}
                    current_product["spec_dict"][spec_key] = value

    if current_product and (current_product.get("model") or current_product.get("name_zh")):
        if current_product.get("spec_dict"):
            current_product["spec_zh"] = "; ".join(
                f"{k}: {v}" for k, v in current_product["spec_dict"].items()
            )
        products.append(current_product)

    if not products:
        return None

    for p in products:
        if "spec_dict" not in p:
            p["spec_dict"] = {}
        if "spec_zh" not in p or not p["spec_zh"]:
            if p.get("spec_dict"):
                p["spec_zh"] = "; ".join(f"{k}: {v}" for k, v in p["spec_dict"].items())
        if not p.get("spec_zh"):
            p["spec_zh"] = ""

    result = pd.DataFrame(products)
    for col in ["model", "name_zh", "spec_zh", "color", "package", "price_rmb", "price_usd", "spec_dict"]:
        if col not in result.columns:
            result[col] = {} if col == "spec_dict" else None

    return result[["model", "name_zh", "spec_zh", "spec_dict", "color", "package", "price_rmb", "price_usd"]]


def extract_excel_images(file_path: str, temp_folder: str = None) -> dict:
    """
    Extract embedded images from Excel file.
    
    FIX: Also extract model/product info from the same row to enable model-based matching.
    Returns: dict with row number as key, but also stores model info for matching.
    """
    if temp_folder is None:
        temp_folder = TEMP_IMAGE_FOLDER
    
    os.makedirs(temp_folder, exist_ok=True)
    
    image_map = {}  # row -> [image_paths]
    model_map = {}  # row -> model_number
    filename = os.path.splitext(os.path.basename(file_path))[0]
    
    if file_path.lower().endswith('.xls') and '.xlsx' not in file_path.lower():
        return image_map
    
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("openpyxl not installed.")
        return image_map
    
    try:
        wb = load_workbook(file_path, data_only=False)
    except Exception:
        return image_map

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        
        # First pass: extract images with their row data
        if hasattr(sheet, '_images') and sheet._images:
            print(f"  Found {len(sheet._images)} embedded images in sheet '{sheet_name}'")
            
            for img_idx, img in enumerate(sheet._images):
                try:
                    row = img.anchor._from.row
                    
                    img_data = img._data()
                    if not img_data:
                        continue

                    img_type = getattr(img, 'format', 'png').lower()
                    ext = '.jpg' if img_type == 'jpeg' else '.png'

                    img_filename = f"{filename}_{sheet_name}_row{row}_{img_idx}{ext}"
                    img_path = os.path.join(temp_folder, img_filename)

                    with open(img_path, 'wb') as f:
                        f.write(img_data)

                    if row not in image_map:
                        image_map[row] = []
                    image_map[row].append(img_path)
                    
                except Exception as e:
                    print(f"  Warning: Failed to extract image {img_idx}: {e}")
                    continue
        
        # Second pass: extract model info from nearby cells
        # Read with openpyxl to get cell values
        try:
            wb_data = load_workbook(file_path, data_only=True)
            sheet_data = wb_data[sheet_name]
            
            for row_num in image_map.keys():
                # Look in columns B, C, D for model info (typical locations)
                for col in ['B', 'C', 'D', 'A']:
                    cell_val = sheet_data[f'{col}{row_num}'].value
                    if cell_val and isinstance(cell_val, str):
                        # Check if it looks like a model number
                        if re.match(r'^[A-Z][0-9/-]+$', cell_val.strip()) or \
                           re.match(r'^[A-Z]{1,3}-[A-Z0-9]+$', cell_val.strip()):
                            model_map[row_num] = cell_val.strip()
                            break
        except Exception as e:
            print(f"  Note: Could not extract model info: {e}")
        
        # Scan for DISPIMG formulas
        dispimg_count = 0
        for row_idx, row in enumerate(sheet.iter_rows()):
            for col_idx, cell in enumerate(row):
                if cell.value and isinstance(cell.value, str):
                    if 'DISPIMG' in cell.value.upper():
                        match = re.search(r'DISPIMG\s*\(\s*["\']([^"\']+)["\']', cell.value, re.IGNORECASE)
                        if match:
                            img_id = match.group(1)
                            dispimg_count += 1
        
        if dispimg_count > 0:
            print(f"  Note: {dispimg_count} DISPIMG formulas found.")
    
    # Store model_map in a global for later matching
    global _last_model_map
    _last_model_map = model_map
    
    return image_map


def parse_excel(file_path: str, exchange_rate: float = DEFAULT_EXCHANGE_RATE, serialize: bool = True) -> Optional[pd.DataFrame]:
    """
    Parse a single Excel file.
    
    Args:
        file_path: Excel文件路径
        exchange_rate: 汇率
        serialize: 是否执行规格参数序列化（默认True）
    
    Returns:
        解析后的DataFrame
    """
    basename = os.path.basename(file_path).lower()
    
    # Handle old .xls format
    if '.xls' in basename and '.xlsx' not in basename:
        if not XLRD_AVAILABLE:
            print(f"Warning: xlrd not installed. Cannot read {file_path}")
            return None
        
        try:
            df = pd.read_excel(file_path, sheet_name=0, header=0, engine='xlrd')
            print(f"  Read .xls file with xlrd engine")
        except Exception as e:
            print(f"Warning: Failed to read {file_path}: {e}")
            return None
    else:
        engine = 'openpyxl'
        ext = os.path.splitext(file_path)[1].lower()
        
        df = None
        encodings = [None, 'utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
        
        for encoding in encodings:
            try:
                kwargs = {'sheet_name': 0, 'header': 0, 'engine': engine}
                if encoding and ext in ['.xls', '.xlsx']:
                    kwargs['encoding'] = encoding
                df = pd.read_excel(file_path, **kwargs)
                if df is not None and len(df) >= MIN_PRODUCT_ROWS:
                    print(f"  Read with encoding: {encoding or 'default'}")
                    break
            except Exception:
                continue
        
        if df is None:
            try:
                df = pd.read_excel(file_path, sheet_name=0, header=None, engine=engine)
            except Exception as e:
                print(f"Warning: Cannot read {file_path}: {e}")
                return None

    if df is None or len(df) < MIN_PRODUCT_ROWS:
        return None

    excel_images = extract_excel_images(file_path)

    # 提前序列化：在normalize之前，使用原始列
    if serialize and len(df.columns) > 6:
        print(f"  Serializing specs before normalization...")
        df_serialized = serialize_specs(df)
        # 如果成功生成spec_text列，用序列化后的df
        if df_serialized is not None and "spec_text_zh" in df_serialized.columns:
            if df_serialized["spec_text_zh"].notna().any():
                df = df_serialized
                print(f"  Generated spec_text columns")

    if _is_key_value_format(df):
        result = _parse_key_value_format(df, exchange_rate)
    else:
        result = _normalize_dataframe(df, exchange_rate)

    if result is not None and len(result) > 0 and excel_images:
        image_paths = []
        
        first_data_row = 0
        for test_idx in range(min(10, len(result))):
            row = result.iloc[test_idx]
            if any(pd.notna(v) for v in row if v):
                first_data_row = test_idx
                break
        
        for idx in result.index:
            excel_row = idx + first_data_row + 1
            
            if excel_row in excel_images:
                image_paths.append("; ".join(excel_images[excel_row]))
            else:
                found = False
                for offset in range(-2, 3):
                    check_row = excel_row + offset
                    if check_row in excel_images:
                        image_paths.append("; ".join(excel_images[check_row]))
                        found = True
                        break
                if not found:
                    image_paths.append(None)
        result["image_path"] = image_paths

    # ========== FIX: Add source metadata ==========
    if result is not None:
        source_filename = os.path.splitext(os.path.basename(file_path))[0]
        result["_source_file"] = source_filename
        
        try:
            xl = pd.ExcelFile(file_path)
            sheet_names = xl.sheet_names
            if sheet_names:
                result["_source_sheet"] = sheet_names[0]
        except:
            pass
    
    return result


# ==================== PRICE CONVERSION ==========
def convert_price_usd(df: pd.DataFrame, exchange_rate: float = DEFAULT_EXCHANGE_RATE) -> pd.DataFrame:
    """
    Convert RMB price to USD.
    
    FIX: 
    - If price is already in USD format (like "0.17/pc" or "0.2"), keep as-is
    - Only convert RMB to USD if price is in RMB format
    
    Check for USD indicators: $, USD, /pc, /box, or small decimal (< 100)
    """
    if df is None or df.empty:
        return df
    
    result = df.copy()
    
    # Check if price_usd already contains valid USD values
    has_original_usd = False
    if "price_usd" in result.columns:
        # Check for USD indicators in the raw values
        usd_indicators = ['$', 'usd', '/pc', '/box', '0.1', '0.2', '0.3', '0.01', '0.05', '0.07']
        for val in result["price_usd"].dropna().head(10):
            val_str = str(val).lower()
            # If has USD indicators, likely already USD
            if any(ind in val_str for ind in usd_indicators):
                has_original_usd = True
                break
            # If numeric value is small (< 100), likely already USD
            try:
                numbers = re.findall(r'[\d.]+', str(val))
                if numbers:
                    num = float(numbers[0])
                    if 0 < num < 100:
                        has_original_usd = True
                        break
            except:
                pass
    
    if has_original_usd:
        # Already USD - just extract the numeric value
        def extract_usd(val):
            if pd.isna(val) or not val:
                return None
            val_str = str(val)
            # Get all numbers and find the first meaningful one
            numbers = re.findall(r'[\d.]+', val_str)
            if numbers:
                try:
                    return round(float(numbers[0]), 2)
                except:
                    return None
            return None
        result["price_usd"] = result["price_usd"].apply(extract_usd)
    elif "price_rmb" in result.columns:
        # Convert from RMB
        def calc_usd(val):
            if pd.isna(val) or not val:
                return None
            numbers = re.findall(r'[\d.]+', str(val))
            if numbers:
                try:
                    # FIX: DIVIDE by exchange rate
                    return round(float(numbers[0]) / exchange_rate, 2)
                except:
                    return None
            return None
        result["price_usd"] = result["price_rmb"].apply(calc_usd)
    
    return result


# ==================== MAIN LOAD FUNCTION ====================

def load_files(input_folder: str, exchange_rate: float = DEFAULT_EXCHANGE_RATE, file_type: str = "all", serialize: bool = True) -> pd.DataFrame:
    """Load all supported files from folder.
    
    Args:
        input_folder: Input folder path
        exchange_rate: Exchange rate for price conversion
        file_type: "all", "excel", "pdf", or "docx"
        serialize: 是否执行规格参数序列化（默认True）
    """
    all_data = []
    
    # Process Excel files
    if file_type in ["all", "excel"]:
        for ext in ["*.xlsx", "*.xls"]:
            for file_path in glob.glob(os.path.join(input_folder, ext)):
                print(f"\nProcessing: {os.path.basename(file_path)}")
                df = parse_excel(file_path, exchange_rate, serialize=serialize)
                
                if df is not None and len(df) > 0:
                    df = convert_price_usd(df, exchange_rate)
                    all_data.append(df)
                    print(f"  Loaded {len(df)} rows")
    
    # Process PDF files
    if file_type in ["all", "pdf"]:
        for ext in ["*.pdf"]:
            for file_path in glob.glob(os.path.join(input_folder, ext)):
                print(f"\nProcessing: {os.path.basename(file_path)}")
                df = parse_pdf(file_path, exchange_rate)
                
                if df is not None and len(df) > 0:
                    all_data.append(df)
                    print(f"  Loaded {len(df)} rows")
    
    # Process Word DOCX files
    if file_type in ["all", "docx"]:
        for ext in ["*.docx"]:
            for file_path in glob.glob(os.path.join(input_folder, ext)):
                print(f"\nProcessing: {os.path.basename(file_path)}")
                df = parse_docx(file_path, exchange_rate)
                
                if df is not None and len(df) > 0:
                    all_data.append(df)
                    print(f"  Loaded {len(df)} rows")
    
    if not all_data:
        return pd.DataFrame()
    
    return pd.concat(all_data, ignore_index=True)


def load_excel_files(input_folder: str, exchange_rate: float = DEFAULT_EXCHANGE_RATE) -> pd.DataFrame:
    """Load all Excel files from folder."""
    all_data = []
    
    for ext in ["*.xlsx", "*.xls"]:
        for file_path in glob.glob(os.path.join(input_folder, ext)):
            print(f"\nProcessing: {os.path.basename(file_path)}")
            df = parse_excel(file_path, exchange_rate)
            
            if df is not None and len(df) > 0:
                # Apply price conversion
                df = convert_price_usd(df, exchange_rate)
                all_data.append(df)
                print(f"  Loaded {len(df)} rows")
    
    if not all_data:
        return pd.DataFrame()
    
    return pd.concat(all_data, ignore_index=True)


def load_single_file(file_path: str, exchange_rate: float = DEFAULT_EXCHANGE_RATE, serialize: bool = True) -> Optional[pd.DataFrame]:
    """Load a single Excel file."""
    print(f"\nProcessing: {os.path.basename(file_path)}")
    df = parse_excel(file_path, exchange_rate, serialize=serialize)
    
    if df is not None and len(df) > 0:
        df = convert_price_usd(df, exchange_rate)
        print(f"  Loaded {len(df)} rows")
        return df
    
    return None


def load_documents(
    input_folder: str,
    exchange_rate: float = DEFAULT_EXCHANGE_RATE,
    use_ocr: bool = False,
    fetch_rate: bool = True,
    file_type: str = "all",
    serialize: bool = True,
) -> pd.DataFrame:
    """Load all supported documents from folder.
    
    Args:
        input_folder: Input folder path
        exchange_rate: Exchange rate
        use_ocr: Use OCR for PDF (not implemented)
        fetch_rate: Fetch real exchange rate
        file_type: "all", "excel", or "pdf"
        serialize: 是否执行规格参数序列化（默认True）
    """
    if fetch_rate:
        get_exchange_rate()
    
    return load_files(input_folder, exchange_rate, file_type, serialize)


# ==================== PLACEHOLDERS ====================

def extract_pdf_images(file_path: str, temp_folder: str = None) -> dict:
    return {"images": [], "unmatched": []}

def match_pdf_images_to_products(pdf_images: dict, df: pd.DataFrame) -> pd.DataFrame:
    return df


def parse_pdf(file_path: str, exchange_rate: float = DEFAULT_EXCHANGE_RATE, use_ocr: bool = False) -> Optional[pd.DataFrame]:
    """Parse PDF files to extract ALL content - enhanced version."""
    import pdfplumber
    
    all_products = []
    basename = os.path.basename(file_path)
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            
            # Extract from tables - get ALL columns
            for table in tables:
                if not table:
                    continue
                
                # Get headers from first row
                headers = None
                if table and len(table) > 0:
                    first_row = [str(c).strip() if c else "" for c in table[0]]
                    # Check if first row looks like headers
                    header_keywords = ['model', 'name', 'type', 'description', 'spec', 'price', 
                                   'specification', 'product', '产品', '型号', '名称', '规格', 
                                   'price', 'color', 'package', 'moq', 'lead', 'certification']
                    if first_row[0] and any(kw in first_row[0].lower() for kw in header_keywords):
                        headers = first_row
                        table = table[1:]  # Skip header row
                
                # Process ALL table rows
                for row in table:
                    if not row or len(row) < 2:
                        continue
                    
                    # Skip empty/header rows
                    if row[0] and isinstance(row[0], str):
                        if row[0].lower() in ['type', 'model', 'models', 'name', 'description', 'total', '合计']:
                            continue
                    
                    # Get model (first column)
                    if not row[0] or not isinstance(row[0], str):
                        continue
                    model = str(row[0]).strip()
                    
                    # Model pattern check - more flexible (e.g., SC-CT001, PS-500, PV-1200)
                    # Matches: 2+ letters + optional dash + numbers/letters
                    # Must have at least 2 chars and contain numbers
                    clean_model = model.replace(' ', '').replace('-', '')
                    
                    # DETECT: If model is just "CE", "PV", "MSDS" etc (short cert/acronym)
                    # and has no numbers, this is likely a spec-sheet, not product table
                    is_cert_or_spec = len(clean_model) <= 5 and not re.search(r'\d', clean_model)
                    
                    if is_cert_or_spec:
                        continue  # Skip these - not real product models
                    
                    # Must contain at least one digit to be a valid product model
                    if not re.search(r'\d', model):
                        continue
                    
                    # Extract ALL columns from the row
                    name = row[1] if len(row) > 1 and row[1] else model
                    spec = []
                    price = None
                    color = None
                    package = None
                    cert = None
                    lead_time = None
                    moq = None
                    
                    # Process each column intelligently
                    for col_idx, cell in enumerate(row):
                        if not cell:
                            continue
                        cell_str = str(cell).strip()
                        if not cell_str:
                            continue
                        
                        # Try to identify column type
                        col_lower = cell_str.lower()
                        
                        # Check if it's a price (has $ or number at end)
                        if col_idx == len(row) - 1:
                            try:
                                price = float(re.findall(r'[\d.]+', cell_str)[0])
                            except:
                                pass
                        
                        # Match patterns for specific fields
                        if 'color' in col_lower or '颜色' in cell_str:
                            color = cell_str
                        elif 'package' in col_lower or '包装' in cell_str:
                            package = cell_str
                        elif 'moq' in col_lower or '最小' in cell_str:
                            try:
                                moq = re.findall(r'\d+', cell_str)[0]
                            except:
                                pass
                        elif 'lead' in col_lower or '交货' in cell_str:
                            lead_time = cell_str
                        elif 'cert' in col_lower or '认证' in cell_str:
                            cert = cell_str
                        elif col_idx >= 2 and col_idx <= 8:
                            # Spec columns
                            spec.append(cell_str)
                    
                    # Build spec string from all spec columns
                    spec_str = "; ".join(spec) if spec else ""
                    
                    product = {
                        "model": model,
                        "name_zh": name,
                        "spec_zh": spec_str,
                        "spec_dict": {},
                        "color": color,
                        "package": package,
                        "price_rmb": None,
                        "price_usd": price,
                        "certification": cert,
                        "lead_time": lead_time,
                        "moq": moq,
                    }
                    all_products.append(product)
            
            # If no products from tables, try text extraction
            if not all_products:
                text = page.extract_text()
                if not text:
                    continue
                
                lines = text.split('\n')
                for line_num, line in enumerate(lines):
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Match: MODEL NAME PRICE (e.g., PS-500 Solar Panel 500W)
                    match = re.match(r'^([A-Z][A-Z0-9\-]+)\s+(.+?)([\d.]+)?$', line)
                    if match:
                        model = match.group(1)
                        name = match.group(2).strip()
                        price = None
                        
                        # MUST have at least one digit in model (e.g., PS-500, not just PV)
                        if not re.search(r'\d', model):
                            continue
                        
                        if match.group(3):
                            try:
                                price = float(match.group(3))
                            except:
                                pass
                        
                        product = {
                            "model": model,
                            "name_zh": name,
                            "spec_zh": "",
                            "spec_dict": {},
                            "color": None,
                            "package": None,
                            "price_rmb": None,
                            "price_usd": price,
                        }
                        all_products.append(product)
    
    # DISABLED: Key-value format creates too many "products" from spec sheets
    # Keeping just in case, but with strict rules
    # if not all_products:
    #     try:
    #         with pdfplumber.open(file_path) as pdf:
    #             # ... only create if filename looks like "catalog"
    #     except:
    #         pass
    
    # SKIP key-value fallback entirely - it misidentifies spec lines as products
    # Instead go directly to FINAL FALLBACK for spec-sheet PDFs
    
    # FINAL FALLBACK: For spec-sheet type PDFs like Portable Power Station
    # Try to extract as SINGLE product with all specs (not multiple)
    if not all_products:
        try:
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                spec_dict = {}
                
                for page in pdf.pages:
                    text = page.extract_text()
                    if not text:
                        continue
                    
                    # Extract ALL text lines as spec details
                    lines = text.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        full_text += line + "\n"
                        
                        # Try key:value format
                        if ':' in line or '—' in line:
                            sep = ':' if ':' in line else '—'
                            parts = line.split(sep, 1)
                            if len(parts) == 2:
                                key = parts[0].strip()
                                value = parts[1].strip()
                                if key and value and len(key) > 1:
                                    spec_dict[key.strip()] = value.strip()
                    
                    if full_text:
                        model_name = os.path.splitext(os.path.basename(file_path))[0].replace(" ", "-")
                        # Create ONE product only with all specs
                        product = {
                            "model": model_name[:50],
                            "name_zh": model_name,
                            "spec_zh": full_text[:2000],  # Limit length
                            "spec_dict": spec_dict if spec_dict else {},
                            "color": None,
                            "package": None,
                            "price_rmb": None,
                            "price_usd": None,
                        }
                        all_products.append(product)
                        break  # Only ONE product
        except Exception:
            pass
    
    if all_products:
        result = pd.DataFrame(all_products)
        for col in ["model", "name_zh", "spec_zh", "spec_dict", "color", "package", 
                   "price_rmb", "price_usd", "certification", "lead_time", "moq"]:
            if col not in result.columns:
                result[col] = None
        print(f"  Parsed {len(result)} products from PDF")
        return result
    
    return None


def parse_docx(file_path: str, exchange_rate: float = DEFAULT_EXCHANGE_RATE) -> Optional[pd.DataFrame]:
    """Parse Word DOCX files to extract product data."""
    from docx import Document
    
    all_products = []
    
    try:
        doc = Document(file_path)
        
        # Extract from tables in Word document
        for table in doc.tables:
            # Get headers from first row
            headers = None
            if table.rows and len(table.rows) > 0:
                first_row = [str(cell.text).strip() if cell.text else "" for cell in table.rows[0].cells]
                header_keywords = ['model', 'name', 'type', 'description', 'spec', 'price', 
                               'specification', 'product', '产品', '型号', '名称', '规格', 'color', 'package']
                if first_row[0] and any(kw in first_row[0].lower() for kw in header_keywords):
                    headers = first_row
                    data_rows = table.rows[1:]
                else:
                    data_rows = table.rows
                
                # Process each data row
                for row in data_rows:
                    cells = [c.text.strip() if c.text else "" for c in row.cells]
                    if not cells or not cells[0]:
                        continue
                    
                    model = cells[0]
                    if not re.match(r'^[A-Z][A-Z0-9\-]{1,20}$', model.replace(' ', '')):
                        continue
                    
                    name = cells[1] if len(cells) > 1 and cells[1] else model
                    spec = "; ".join(cells[2:]) if len(cells) > 2 else ""
                    price = None
                    if len(cells) > 1:
                        try:
                            price = float(re.findall(r'[\d.]+', cells[-1])[0])
                        except:
                            pass
                    
                    product = {
                        "model": model,
                        "name_zh": name,
                        "spec_zh": spec,
                        "spec_dict": {},
                        "color": None,
                        "package": None,
                        "price_rmb": None,
                        "price_usd": price,
                    }
                    all_products.append(product)
        
        # If no products from tables, try paragraphs
        if not all_products:
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Match: MODEL NAME PRICE
                match = re.match(r'^([A-Z][A-Z0-9\-]+)\s+(.+?)([\d.]+)?$', text)
                if match:
                    model = match.group(1)
                    name = match.group(2).strip()
                    price = None
                    if match.group(3):
                        try:
                            price = float(match.group(3))
                        except:
                            pass
                    
                    product = {
                        "model": model,
                        "name_zh": name,
                        "spec_zh": "",
                        "spec_dict": {},
                        "color": None,
                        "package": None,
                        "price_rmb": None,
                        "price_usd": price,
                    }
                    all_products.append(product)
    
    except Exception as e:
        print(f"  Warning: Failed to parse DOCX: {e}")
    
    if all_products:
        result = pd.DataFrame(all_products)
        for col in ["model", "name_zh", "spec_zh", "spec_dict", "color", "package", "price_rmb", "price_usd"]:
            if col not in result.columns:
                result[col] = None
        print(f"  Parsed {len(result)} products from DOCX")
        return result
    
    return None


if __name__ == "__main__":
    import sys
    sys.path.insert(0, ".")
    print("Loading data...")
    df = load_excel_files("./data")
    print(f"\nTotal: {len(df)} products")